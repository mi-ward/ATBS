import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH


doc = docx.Document("invites.docx")
guests = open('guests.txt')

for guest in guests:                
    paragraph1 = doc.add_paragraph("It would be a pleasure to have the company of")
    paragraph1.style = 'goofy'
    paragraph1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    paragraph2 = doc.add_paragraph(guest.strip())
    paragraph2.style = 'fancy'
    paragraph2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    paragraph3 = doc.add_paragraph("at 11010 Memory Lane on the Evening of")
    paragraph3.style = 'goofy'
    paragraph3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    paragraph4 = doc.add_paragraph("April 1st")
    paragraph4.style = 'fancy'
    paragraph4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    paragraph5 = doc.add_paragraph("at 7 0'clock")
    paragraph5.style = 'goofy'
    paragraph5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    
doc.save("completedInvites.docx")